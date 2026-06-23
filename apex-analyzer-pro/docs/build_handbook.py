# -*- coding: utf-8 -*-
"""
Generates the Apex Analyzer Pro Engineering Handbook & User Guide as a .docx.
Run:  python docs/build_handbook.py
Output: C:/Users/DARSHAN/Desktop/Apex_Analyzer_Pro_Handbook.docx
"""
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\DARSHAN\Desktop\Apex_Analyzer_Pro_Handbook.docx"
ACCENT = RGBColor(0x21, 0x3F, 0x9E)
DARK = RGBColor(0x13, 0x18, 0x22)
SLATE = RGBColor(0x47, 0x53, 0x61)

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, sz, col in [(1, 17, ACCENT), (2, 13.5, ACCENT), (3, 11.5, DARK)]:
    st = doc.styles['Heading %d' % lvl]
    st.font.name = 'Calibri'
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = col

# ---- page geometry (US Letter, 1in margins) ----
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(sec, m, Inches(1))


# NOTE on ordering: inside <w:pPr> the schema requires pBdr BEFORE shd, and both
# BEFORE spacing/ind/jc/rPr. We insert at the front (shd first, then pBdr) so the
# final order is [pBdr, shd, ...]; python-docx then inserts spacing after them.
def _shade_para(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.insert(0, shd)


def _border_para(p, color='C9D2E3'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6'); e.set(qn('w:space'), '6'); e.set(qn('w:color'), color)
        pBdr.append(e)
    pPr.insert(0, pBdr)


def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def H1(t):
    doc.add_heading(t, level=1)

def H2(t):
    doc.add_heading(t, level=2)

def H3(t):
    doc.add_heading(t, level=3)

def P(t, bold=False, italic=False, size=None, color=None, lead=None):
    p = doc.add_paragraph()
    if lead:
        r = p.add_run(lead + '  ')
        r.bold = True
    r = p.add_run(t)
    r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def BUL(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ' '); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)

def NUM(items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        if isinstance(it, tuple):
            r = p.add_run(it[0] + ' '); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)

def CODE(code):
    p = doc.add_paragraph()
    _shade_para(p, 'F3F5F9'); _border_para(p)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    lines = code.strip('\n').split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line if line != '' else ' ')
        run.font.name = 'Consolas'; run.font.size = Pt(8.5)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), 'Consolas'); rfonts.set(qn('w:hAnsi'), 'Consolas')
        if i < len(lines) - 1:
            run.add_break()
    return p

CONTENT_W = 6.5  # US Letter, 1in margins -> 6.5in printable width


def _apply_widths(t, widths_in):
    """Pin column widths reliably: rewrite the tblGrid AND every cell width.
    With fixed layout (table.autofit = False) Word honours these literally."""
    grid = t._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for gc in list(grid.findall(qn('w:gridCol'))):
            grid.remove(gc)
        for w in widths_in:
            gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(round(w * 1440))))
            grid.append(gc)
    for row in t.rows:
        for c_i, w in enumerate(widths_in):
            row.cells[c_i].width = Inches(w)


def TBL(headers, rows, widths=None, header_fill='2B3445'):
    n = len(headers)
    # Normalize requested proportions to the printable width so nothing overflows.
    if widths:
        s = float(sum(widths))
        widths = [w / s * CONTENT_W for w in widths]
    else:
        widths = [CONTENT_W / n] * n
    t = doc.add_table(rows=1, cols=n)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False  # fixed layout
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[i], header_fill)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
    _apply_widths(t, widths)
    return t

def PB():
    doc.add_page_break()

def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Right-click here and choose "Update Field" to generate the contents.'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    for x in (b, instr, sep, t, e):
        run._r.append(x)

def add_footer():
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Apex Analyzer Pro - Engineering Handbook & User Guide    |    Page ')
    run = p.add_run()
    for typ, txt in (('begin', None), ('instr', 'PAGE'), ('end', None)):
        if typ == 'instr':
            el = OxmlElement('w:instrText'); el.set(qn('xml:space'), 'preserve'); el.text = ' PAGE '
        else:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), typ)
        run._r.append(el)
    for r in p.runs:
        r.font.size = Pt(8); r.font.color.rgb = SLATE


# =====================================================================
# TITLE PAGE
# =====================================================================
for _ in range(3):
    doc.add_paragraph()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('APEX ANALYZER PRO'); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = DARK
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Engineering Handbook & User Guide'); r.font.size = Pt(18); r.font.color.rgb = ACCENT
sp2 = doc.add_paragraph(); sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp2.add_run('Competitor Intelligence & Pivot Workbench (Chrome Extension, Manifest V3)')
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = SLATE
for _ in range(8):
    doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version 1.0.0\n').bold = True
meta.add_run('Build target: Chrome / Chromium MV3\n')
meta.add_run('Document generated: %s\n' % datetime.date.today().isoformat())
meta.add_run('Audience: developers, DevOps engineers, and power users')
for r in meta.runs:
    r.font.size = Pt(11); r.font.color.rgb = SLATE
PB()

# =====================================================================
# TOC
# =====================================================================
H1('Table of Contents')
add_toc()
PB()

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
H1('1. Introduction')

H2('1.1 What Apex Analyzer Pro is')
P('Apex Analyzer Pro is a self-contained, offline-first Chrome extension that turns a scraped '
  'competitor price/specification export (an Excel or CSV file) into an interactive competitive-'
  'intelligence workbench. Its centrepiece is a configurable pivot table with cell-level drill-down; '
  'around that sit a market overview, a specification-bucket comparison matrix, time-series trends, a '
  'two-brand head-to-head comparator, a small SQL workspace, and an AI/alerting layer. The reference '
  'domain is the Indian mattress market (brands such as Sleepwell, Kurlon, Wakefit, Duroflex), but the '
  'engine is domain-agnostic and works on any comparably shaped dataset.')

H2('1.2 Who this handbook is for')
BUL([
    ('Developers', 'who must build the tool from scratch, maintain it, or extend it. Parts 5, 6, 8 and 11 are for you.'),
    ('DevOps engineers', 'who must set up the build/test/release pipeline and distribution. Parts 8, 9 and 10 are for you.'),
    ('Power users / analysts', 'who operate the tool day to day. Part 4 (User Guide) is for you.'),
])

H2('1.3 Design principles')
NUM([
    ('Engine-first, DOM-free core.', 'All data logic lives in engine.js, which has zero browser dependencies and is unit-testable in Node. The UI (popup.js) only renders what the engine computes.'),
    ('Ground-truth correctness.', 'Every aggregation is verified to match the source workbook to the decimal. Numbers must be trustworthy before they are pretty.'),
    ('No build step.', 'Plain ES5-compatible JavaScript, no transpiler/bundler. Load-unpacked and reload; what you write is what runs. This keeps the tool auditable and dependency-light.'),
    ('Offline-first and private.', 'Nothing leaves the machine. The optional local LLM (Ollama) and the optional outbound webhook are the only network paths, and both are user-initiated.'),
    ('No emoji or icon glyphs.', 'Every affordance is typographic or geometric (CSS-drawn). The only non-ASCII glyph used is the rupee currency sign.'),
])

H2('1.4 Honest scope - "the tool is not perfect"')
P('This handbook is candid about limitations. The tool is production-useful but deliberately bounded: '
  'the SQL engine is a practical subset (no JOINs/sub-queries), the pivot supports up to two row and two '
  'column levels, all processing is single-threaded and in-memory, and there is no automated test of the '
  'DOM layer beyond scripted browser smoke checks. Section 13 lists every known limitation and the '
  'roadmap to address it. Read it before promising capabilities to stakeholders.')

PB()

# =====================================================================
# 2. ARCHITECTURE
# =====================================================================
H1('2. System Architecture')

H2('2.1 High-level shape')
P('Apex Analyzer Pro is a two-layer single-page application packaged as an MV3 extension:')
BUL([
    ('Pure analytics engine (engine.js).', 'Header mapping, normalization, cleaning, the pivot engine with a drill-down index, KPI/summary builders, the A-vs-B comparator, the alert engine, and a mini SQL parser. No DOM, no globals beyond one namespace. UMD-exported so it runs identically in the browser and in Node tests.'),
    ('UI controller (popup.js).', 'Owns all DOM, state, event wiring, rendering, drill-down drawer, exports, and the notification layer. It calls into the engine and never re-implements analytics.'),
    ('Markup & styling (popup.html, popup.css).', 'A static shell with seven panes and a drill-down drawer; a token-based design system with no icon assets.'),
    ('Extension plumbing (manifest.json, background.js).', 'MV3 manifest plus a service worker that opens the workbench in a full browser tab.'),
    ('Third-party (xlsx.mini.min.js).', 'SheetJS community build for reading/writing .xlsx and .csv. Trusted boundary - the engine never parses spreadsheets itself.'),
])

H2('2.2 Runtime data flow')
CODE(
"User picks file\n"
"  -> handleFile()           (popup.js)  FileReader -> ArrayBuffer\n"
"  -> XLSX.read()            (SheetJS)    workbook object\n"
"  -> chooseSheet()          (popup.js)  auto-detect 'Consolidated Data'\n"
"  -> XLSX.utils.sheet_to_json()          array of raw row objects\n"
"  -> ingestRows()           (popup.js)  orchestrates the engine:\n"
"        ApexEngine.normalizeRows(raw)    -> canonical records\n"
"        ApexEngine.cleanRows(records)    -> { clean, dropped }\n"
"  -> refreshAll()           (popup.js)  re-render every pane from STATE.clean\n"
"\n"
"Pivot interaction:\n"
"  currentPivotConfig() -> ApexEngine.buildPivot(clean, cfg) -> render -> click cell\n"
"  -> onDrill() reads pivot.drill[cellId] (row indices) -> openDrawer()"
)

H2('2.3 Technology stack and rationale')
TBL(['Concern', 'Choice', 'Why'],
    [['Platform', 'Chrome MV3 extension', 'Runs locally, no server, distributable to a team'],
     ['Language', 'Vanilla JS (ES5-compatible)', 'No build step; auditable; broad engine compatibility (also runs in Node for tests)'],
     ['Spreadsheets', 'SheetJS (xlsx.mini)', 'Battle-tested .xlsx/.csv read & write in the browser'],
     ['Charts', 'Pure CSS bars / HTML tables', 'Zero chart-library dependency; honours the no-icon rule'],
     ['AI', 'Ollama (localhost) optional', 'Local LLM; transparent statistical fallback when absent'],
     ['Storage', 'chrome.storage.local + shim', 'Persists settings; shim lets it run as a plain page too'],
     ['Tests', 'Node (assertions) + Python (fixtures)', 'Engine is DOM-free, so it tests headlessly against real data']],
    widths=[1.6, 2.4, 5.0])

H2('2.4 Repository layout')
TBL(['Path', 'Purpose'],
    [['manifest.json', 'MV3 manifest: name, permissions, action, background worker'],
     ['background.js', 'Service worker; opens popup.html in a tab on toolbar click'],
     ['popup.html', 'Workbench markup: top bar, sidebar, seven panes, drill drawer'],
     ['popup.css', 'Design system (tokens, components); no icon assets'],
     ['engine.js', 'Pure analytics engine (UMD); the heart of the tool'],
     ['popup.js', 'UI controller; wires the engine to the DOM'],
     ['xlsx.mini.min.js', 'SheetJS spreadsheet read/write library'],
     ['icon16/48/128.png, favicon.ico', 'Toolbar/app icons (extension chrome, not in-UI)'],
     ['test/backtest.js', 'Node harness: 59 assertions vs workbook ground truth'],
     ['test/ingest_test.js', 'Node harness: real .xlsx end-to-end ingestion check'],
     ['docs/build_handbook.py', 'This document generator'],
     ['README.md', 'Quick-start and feature summary']],
    widths=[2.6, 6.4])

PB()

# =====================================================================
# 3. DATA MODEL
# =====================================================================
H1('3. Data Model & Semantics')

H2('3.1 Source data expectations')
P('The tool is built around a wide "Consolidated Data" sheet: one row per scraped listing per crawl '
  'date. It thrives on standardized columns when present and degrades gracefully to title-parsing when '
  'they are absent. The reference workbook carries 55 columns and ~33k rows across multiple daily '
  'snapshots. Only a subset is needed; the rest is ignored.')

H2('3.2 Header alias resolution (two-phase mapping)')
P('Source headers vary ("Brand Name" vs "Brand (Std)", "Active Price" vs "price"). buildHeaderMap '
  'resolves each canonical field from a priority-ordered alias list. Resolution is two global phases so '
  'a precise match always beats a loose one, and no column is ever bound by two fields:')
NUM([
    ('Phase 1 - exact match.', 'For every field, in alias order, claim the first header whose cleaned key equals an alias. Claimed headers are marked used.'),
    ('Phase 2 - substring match.', 'For still-unresolved fields, claim the first unclaimed header whose cleaned key contains an alias.'),
])
P('This prevents the classic collision where a loose "height" alias on the thickness field would steal '
  'the "Height (in)" column before the height field could claim it exactly. Standardized columns are '
  'listed first in the alias table so they win.')
P('Header keys are normalized by cleanKey(): lower-cased, with every non-alphanumeric character '
  'stripped. So "Wow/Best Price (Hist)" becomes "wowbestpricehist".')

H2('3.3 Canonical record schema')
P('normalizeRows() converts each raw row into a canonical record. Downstream code only ever sees these '
  'fields:')
TBL(['Field', 'Type', 'Source / meaning'],
    [['idx', 'number', 'Stable original-row index (drill-down key)'],
     ['platform', 'string', 'Amazon / Flipkart / ...'],
     ['brand', 'string', 'Standardized brand (Brand (Std) preferred; else title-cased raw)'],
     ['brandRaw', 'string', 'Brand exactly as scraped'],
     ['title', 'string', 'Listing title'],
     ['url, pid', 'string', 'Product link and ASIN/FSN id'],
     ['price', 'number|null', 'Active price (>0) else effective price'],
     ['effPrice, mrp, wow, emi', 'number|null', 'Effective / list / best-wow / EMI price'],
     ['discount', 'number|null', 'Percent off'],
     ['rating', 'number|null', 'Stars 0-5'],
     ['reviews', 'number|null', 'Review count'],
     ['bsr', 'number|null', 'Bestseller rank (lower = better)'],
     ['type', 'string', 'Product type (std column or title fallback)'],
     ['size', 'string', 'Mattress size'],
     ['thickness', 'string', 'Thickness in inches (as string)'],
     ['dim', 'string', 'Dimensions L x B x H'],
     ['length, breadth, heightDim', 'number|null', 'Parsed L / B / H in inches'],
     ['availability, inStock', 'string', 'Raw availability and resolved In/Out/Unknown'],
     ['prime', 'string', 'Yes / No (prime / assured)'],
     ['seller, warranty, currency', 'string', 'Fulfilment, warranty text, currency'],
     ['scrapeDate', 'string', 'YYYY-MM-DD (or empty)'],
     ['recordStatus', 'string', 'ok / partial / failed / blocked'],
     ['raw', 'object', 'The original row, for drill-down detail/export']],
    widths=[2.1, 1.5, 5.4])

H2('3.4 Cleaning rules')
P('cleanRows() is deliberately conservative to preserve parity with the source workbook. It removes only '
  'rows that are explicit scraper failures (recordStatus failed/blocked) or completely empty (no brand, '
  'no title, no URL, no price). Borderline rows are kept; a row with a blank scrape date simply falls '
  'under the "(blank)" date bucket rather than polluting a dated snapshot. In the reference workbook this '
  'drops zero rows, so a snapshot count of 8,087 matches the workbook exactly.')

H2('3.5 Dimensions catalogue')
P('Dimensions are the things you can group/filter by. Each has a key, a label, and a get(record) '
  'accessor that returns its display value (with "(blank)" for missing values).')
TBL(['Key', 'Label', 'Notes'],
    [['scrapeDate', 'Scrape Date', 'YYYY-MM-DD'],
     ['platform', 'Platform', ''],
     ['brand', 'Brand', 'Standardized'],
     ['type', 'Product Type', ''],
     ['size', 'Mattress Size', ''],
     ['thickness', 'Thickness (in)', 'Formatted "6 in"'],
     ['dim', 'Dimensions', 'L x B x H token'],
     ['length / breadth / height', 'Length/Breadth/Height (in)', 'Independent numeric axes'],
     ['inStock', 'Availability', 'In Stock / Out of Stock / Unknown'],
     ['prime', 'Prime / Assured', ''],
     ['seller', 'Seller', ''],
     ['warranty', 'Warranty', '']],
    widths=[2.2, 2.3, 4.5])

H2('3.6 Measures catalogue and null semantics')
P('Measures are the numbers you aggregate. Each measure get() returns a number or null, applying its '
  'own validity rule so unscraped values do not distort averages:')
BUL([
    ('Positive-only (null if <= 0):', 'price, effPrice, wow, mrp, rating, bsr. A zero here means "not scraped", so it is excluded from averages, sums, min and max.'),
    ('Non-negative (null if < 0):', 'discount, reviews. Zero is a legitimate value (no discount, no reviews) and is kept.'),
])
P('Aggregations: Average, Sum, Minimum, Maximum, Median operate over the non-null measure values in a '
  'cell. Count is special - it counts every row assigned to the cell regardless of measure nullity, '
  'which is why a brand can show 4,016 models but an average computed over slightly fewer priced rows.')

H2('3.7 Workbook-parity rules (memorize these)')
TBL(['Quantity', 'Definition used', 'Reason'],
    [['Average price', 'mean of prices > 0', 'Excel AVERAGE ignores blanks; 0 means unscraped'],
     ['% in stock', 'in-stock / TOTAL models', 'Matches Type Master (e.g. Coir 483/490 = 0.98571)'],
     ['Models (count)', 'all rows in the group', 'Matches Type Master / A-vs-B counts exactly'],
     ['Brand identity', 'Brand (Std) merges aliases', 'KURLON + Kurl-On -> "Kurlon" (9,903 rows)'],
     ['Snapshot', 'latest non-blank scrape date', '2026-06-19 in the reference data']],
    widths=[1.8, 3.0, 4.2])

PB()

# =====================================================================
# 4. USER GUIDE
# =====================================================================
H1('4. User Guide')

H2('4.1 Installation')
NUM([
    'Open chrome://extensions.',
    'Enable Developer mode (top-right toggle).',
    'Click "Load unpacked" and select the apex-analyzer-pro folder.',
    'Click the Apex Analyzer Pro toolbar icon - the workbench opens in a new tab.',
    'After any code change, return to chrome://extensions and click the reload (circular arrow) icon on the card.',
])

H2('4.2 Loading data')
P('Click "Load Data File" in the top bar and choose a .xlsx, .xls, or .csv. The "Consolidated Data" '
  'sheet is auto-detected (or the largest data sheet). If your workbook has several sheets, use the '
  '"Sheet" button to cycle to the right one. The top bar then shows row count, sheet name, brand count '
  'and date count. All seven panes refresh from the loaded data.')

H2('4.3 Overview')
P('Headline metrics for one snapshot. Controls: Snapshot date, Platform and Product Type filters. '
  'Shows eight KPI tiles (listings, brands, platforms, crawl dates, average price/rating/discount, '
  'in-stock %), a sortable Brand Leaderboard (now including Bestseller-Rank stats: Average Rank, Top '
  'Rank and Lowest Rank), an average-price bar chart, and a Product-Type table that reproduces the '
  'workbook Type Master sheet. Export the type table to Excel with one click.')

H2('4.4 Pivot Explorer (the centrepiece)')
P('Build any cross-tab and drill into it. Configure: Rows level 1 and 2, Columns level 1 and 2, a Value '
  '(measure) and an Aggregation. Below that, eight filter dropdowns (Scrape Date, Platform, Brand, '
  'Product Type, Mattress Size, Length, Breadth, Height) scope the data. Toolbar: Heatmap toggle, Row '
  'subtotals toggle, Swap rows/columns, Reset, and Export pivot to Excel.')
BUL([
    ('Drill-down:', 'click any value cell to open a drawer listing the exact underlying listings (brand, title, platform, prices, rating, reviews, stock, date and a product link). The drawer is sortable and exportable.'),
    ('Heatmap:', 'cell background shades from light to dark across the value range for fast visual scanning.'),
    ('Subtotals & totals:', 'two-level row pivots emit per-group subtotals; every pivot has a grand-total row and a total column.'),
])

H2('4.5 Competitive Matrix')
P('Lock a specification bucket (Scrape Date, Product Type, Size, Thickness, Platform, Length, Breadth, '
  'Height) and compare every brand listing in that exact bucket side by side, sorted by price, with a '
  '"vs market average" delta column (green below average, red above) and a clickable "Open listing" link '
  'to each product page. Summary chips show listing count, brand count, market average and the cheapest '
  'listing. Exportable.')

H2('4.6 Trends')
P('Two sections. "Measure over time" pivots a chosen measure across crawl dates, broken down by any '
  'dimension, with a first-to-last delta column; scope it with the Platform/Type/Size/L/B/H filters. '
  '"Snapshot A vs B" matches products by ID across two crawl dates and shows per-product price and '
  'rating movement. Both exportable.')

H2('4.7 A vs B Compare')
P('Head-to-head metrics for two brands across a chosen Type, Size, Platform, Thickness, Length, Breadth, '
  'Height and Snapshot. Twelve metrics (model count, in/out of stock, % in stock, prime SKUs, average '
  'price/wow/discount/rating, lowest price, total reviews, average rank) with an A - B difference column, '
  'plus a "Top brands in competition" panel respecting the same filters. Mirrors the workbook A-vs-B sheet.')

H2('4.8 SQL Workspace')
P('Run ANSI-style SELECT queries against the normalized dataset (table name is ?). Supported grammar: '
  'SELECT fields | *, WHERE with AND/OR/LIKE and comparison/arithmetic operators, GROUP BY with '
  'COUNT(*)/AVG/SUM/MIN/MAX, ORDER BY ... ASC|DESC, and LIMIT. Six sample queries are provided; the '
  'schema panel lists every queryable column. Results are sortable and exportable.')
CODE("SELECT brand, COUNT(*) AS models, AVG(price) AS avg_price\n"
     "FROM ? WHERE scrapeDate = '2026-06-19' GROUP BY brand ORDER BY models DESC")

H2('4.9 AI Insights')
P('Generates a competitive briefing. If a local Ollama server is running at localhost:11434 it uses your '
  'chosen model; otherwise it falls back to a transparent statistical engine that computes every figure '
  'directly from your data (price positioning, identical-spec undercuts, rating watch, recommended '
  'actions). Nothing is sent off-machine in either case except the local Ollama call.')

H2('4.10 Alert Center')
P('Scans the data for what is wrong from the anchor brand\'s point of view and ranks issues by severity. '
  'Five detectors:')
TBL(['Detector', 'What it flags', 'Severity logic'],
    [['Undercut', 'Competitor cheaper on an identical spec', '>=15% high, >=7% medium, else low (>=1% floor)'],
     ['Rating', 'Anchor listing below the rating floor', '< 3.8 high, else medium'],
     ['Stock', 'Anchor listing out of stock', 'medium'],
     ['Competitor price drop', 'Competitor cut price between snapshots', '>=12% high, else medium (>=5% floor)'],
     ['Discount gap', 'Anchor much less promotional than market', 'low (>=10pt gap)']],
    widths=[1.8, 3.4, 3.8])
P('Tunable thresholds: Rating floor and Undercut-high %. Output is a severity summary, a list of action '
  'items (deduplicated, non-low), and a sortable alert table. The scan re-runs when the anchor brand or '
  'thresholds change.')

H2('4.11 Email & notifications')
P('Push the AI report or the critical alerts + action items to your team. Two channels:')
BUL([
    ('Email client (mailto):', 'zero setup. Opens your default mail app with recipients, subject and a plain-text digest pre-filled. Body is capped (~1800 chars) to respect mailto limits.'),
    ('Webhook (POST JSON):', 'for automated delivery. Paste an endpoint URL (Zapier, Make, Google Apps Script, or an SMTP relay). The tool POSTs { to, subject, html, text, alerts, actionItems, generatedAt } and the endpoint forwards the email. The extension requests host permission for that origin at send time and falls back to mailto on any failure.'),
])
P('Recipients, channel and webhook URL are saved on the machine (chrome.storage.local). Note: a browser '
  'extension cannot open raw SMTP sockets, so true unattended email requires the webhook relay - this is '
  'the correct and secure pattern, not a limitation to work around.')

H2('4.12 Exports')
P('Every analytical surface exports to .xlsx via SheetJS: the type table, the pivot (flattened with row '
  'labels, all leaf columns and totals), matrix rows, snapshot A/B comparison, SQL results, and any '
  'drill-down slice. Files download with descriptive names.')

H2('4.13 Troubleshooting')
TBL(['Symptom', 'Likely cause', 'Fix'],
    [['"No dataset loaded"', 'File not parsed / wrong sheet', 'Use the Sheet button; confirm the sheet has a header row and data'],
     ['A filter shows no rows', 'Over-constrained bucket', 'Reset filters to All; widen one axis'],
     ['Pivot all "(blank)"', 'Column not mapped', 'Check your headers against the alias table (Appendix A)'],
     ['AI says "no local LLM"', 'Ollama not running', 'Start "ollama serve" and pull a model, or use the statistical fallback'],
     ['Webhook not delivering', 'Permission denied / bad URL', 'Approve the host permission prompt; verify the endpoint; otherwise use mailto'],
     ['Numbers differ from workbook', 'Different cleaning/measure rule', 'Review Section 3.7 parity rules']],
    widths=[2.2, 2.6, 4.2])

PB()

# =====================================================================
# 5. ENGINE.JS REFERENCE
# =====================================================================
H1('5. Functional Reference - engine.js')
P('engine.js is wrapped in an IIFE and UMD-exported. Its entire public surface is the ApexEngine object '
  '(also module.exports in Node). Below, every function is documented with its role and the code path it '
  'sits on. Line numbers are approximate and may drift; search by function name.')

H2('5.1 Public surface (ApexEngine)')
CODE("ApexEngine = {\n"
     "  toNum, toISODate, classifyFromTitle, buildHeaderMap, normalizeRows, cleanRows,\n"
     "  DIMENSIONS, MEASURES, AGGS, dimByKey, measureByKey, aggregate, applyFilters, buildPivot,\n"
     "  latestSnapshot, distinctDates, distinctValues, computeKpis, groupSummary, compareAvsB,\n"
     "  rowMatchesOpts, buildAlerts, runSql, smartSort\n"
     "}")

H2('5.2 Primitive helpers')
TBL(['Function', 'Purpose & notes'],
    [['cleanKey(s)', 'Lower-case + strip all non-alphanumerics. Header-matching key. "Wow/Best Price (Hist)" -> "wowbestpricehist".'],
     ['toNum(v)', 'Parse a number out of a messy cell ("Rs 12,082" -> 12082). Returns null when nothing numeric is present. Single source of numeric truth.'],
     ['titleCaseBrand(b)', 'Title-cases a raw brand token (tokens <=2 chars upper-cased). Used only when Brand (Std) is absent.'],
     ['toISODate(v)', 'Normalizes a date cell to YYYY-MM-DD (handles Date objects, ISO, d/m/Y, free text). Empty string if not derivable.'],
     ['pad2(n)', 'Zero-pads to two digits for date assembly.']],
    widths=[2.2, 6.8])

H2('5.3 classifyFromTitle(title)')
P('Fallback classifier used ONLY when a row lacks standardized type/size columns. Keyword-matches the '
  'title to infer { type, size, thickness, dim }. Keeps the tool useful on simpler exports; never '
  'overrides a real standardized column.')

H2('5.4 FIELD_ALIASES and buildHeaderMap(headers)')
P('FIELD_ALIASES maps each canonical field to a priority-ordered alias list (standardized names first). '
  'buildHeaderMap returns { field -> originalHeader } using the two-phase algorithm (Section 3.2): all '
  'exact matches claim first, then substring matches over unclaimed headers. The used-tracking guarantees '
  'one header binds to one field, eliminating the Height/Thickness collision class of bug. This function '
  'is the single most correctness-critical mapping step; change it only with the regression tests open.')

H2('5.5 normalizeRows(rawRows) and field resolvers')
P('Builds the header map once, then maps each raw row to a canonical record via the internal pick(row, '
  'field) helper. Brand resolution prefers Brand (Std), else title-cases the raw brand. price prefers a '
  'positive Active Price, else Effective Price. Supporting resolvers:')
BUL([
    ('cleanStr(v)', '- null-safe trim to string.'),
    ('resolveStock(flag, avail)', '- collapses In-Stock flag + availability text into In Stock / Out of Stock / Unknown.'),
    ('resolveBool(v)', '- yes/true/1 -> "Yes"; no/false/0 -> "No".'),
    ('normWarranty(v)', '- light casing normalization of warranty text.'),
])

H2('5.6 cleanRows(records)')
P('Returns { clean, dropped, total }. Conservative: drops only failed/blocked records and completely '
  'empty rows (Section 3.4). Preserves the historical wow-price block (blank scrape date but real prices).')

H2('5.7 Catalogues and lookups')
P('DIMENSIONS, MEASURES, AGGS are the arrays the UI introspects to build every selector. dimByKey / '
  'measureByKey resolve by key. posOrNull (>0 else null) and nonNegOrNull (>=0 else null) implement the '
  'measure validity rules in Section 3.6.')

H2('5.8 aggregate(values, agg, rowCount)')
P('Pure reducer over an array of numbers. agg in {avg,sum,min,max,median}. The special "count" returns '
  'rowCount (all rows in the cell), independent of the values array - this is why count and average can '
  'have different denominators.')

H2('5.9 applyFilters(data, filters)')
P('filters is { dimKey: [allowedValue, ...] }. A row passes only if, for every active filter dimension, '
  'dim.get(row) is in the allowed list. Comparing through dim.get keeps option values and row values on '
  'the same representation (this is why the matrix/trends filters are robust).')

H2('5.10 buildPivot(data, cfg) - the core algorithm')
P('cfg = { rows:[dimKey..], cols:[dimKey..], measure, agg, filters, includeRowSubtotals }. Supports up '
  'to two row and two column levels. Steps:')
NUM([
    'Resolve row/column dimension objects and the measure; apply filters.',
    'distinctVals() computes sorted distinct values per axis field (numeric-aware via smartSort).',
    'Build leafCols (the flat list of data columns) and colTree (the grouped header structure) - one or two column levels, plus an implicit grand-total column.',
    'For each row group (and, for two-level rows, each sub-group and a per-group subtotal), makeRow() computes one cell per leaf column via cellFor().',
    'cellFor() collects the row indices for the (row-path, col-path) intersection, extracts non-null measure values, calls aggregate(), and records the index list in the drill map under a fresh cellId.',
    'Emit a grand-total row across everything.',
])
P('Returns { rowFields, colFields, measure, agg, colTree, leafCols, rows, drill, filteredCount }. The '
  'drill object maps every cellId to its array of record indices - this is what powers cell drill-down. '
  'Invariant verified by tests: the sum of a row\'s leaf-cell counts equals its total-cell count, and '
  'every drill list length equals its cell count.')

H2('5.11 smartSort(a, b)')
P('Natural-ish comparator: "(blank)" sorts last, numeric-leading strings sort numerically, otherwise '
  'locale string compare. Used for all axis ordering.')

H2('5.12 KPI & summary builders')
TBL(['Function', 'Returns / purpose'],
    [['latestSnapshot(data)', 'Most recent non-blank scrape date.'],
     ['distinctDates(data)', 'Sorted unique scrape dates.'],
     ['distinctValues(data, dimKey)', 'Sorted unique dim.get values (drives selector option lists).'],
     ['mean(arr)', 'Null-safe average.'],
     ['computeKpis(data)', 'rows, brands, platforms, dates, avgPrice, avgRating, avgDiscount, inStockPct.'],
     ['groupSummary(data, dimKey)', 'Per-group: count, in/out stock, %inStock (in/total), prime, avgPrice/Wow, lowest, avgDiscount/Rating, totalReviews, and ranking stats avgRank/bestRank/worstRank/rankedCount (Bestseller Rank). Powers the brand leaderboard and Type Master table.']],
    widths=[2.6, 6.4])

H2('5.13 rowMatchesOpts, OPT_DIM, compareAvsB')
P('OPT_DIM maps comparison-filter option keys to dimension keys. rowMatchesOpts(r, opts) returns true '
  'unless some active opt differs from dim.get(r) - comparing through the getter is what makes the '
  '"6 in" thickness option and the "(blank)" option match correctly. compareAvsB(data, A, B, opts) '
  'filters each brand\'s subset through rowMatchesOpts and returns the twelve-metric block for A and B.')

H2('5.14 buildAlerts(data, anchorBrand, opts)')
P('The alerting engine. opts: ratingFloor (default 4.2), undercutHighPct (15), undercutMedPct (7), '
  'maxPerCategory (25). Runs five detectors (Section 4.10), assigns severity, sorts high-first, '
  'deduplicates action items (non-low only, capped at 12) and returns { alerts, counts, total, anchor, '
  'actionItems, dates }. Key correctness details: undercuts apply a >=1% floor (no "-0%" noise); '
  'competitor price-drop matching requires a stable pid or url (never title) and keeps the cheapest row '
  'per key deterministically.')

H2('5.15 The mini SQL engine')
TBL(['Function', 'Role'],
    [['tokenize(str)', 'Lexer: strings, numbers, identifiers, operators, parens; AND/OR/LIKE keywords.'],
     ['PREC', 'Operator precedence table for the shunting-yard parser.'],
     ['shunt(tokens)', 'Shunting-yard -> Reverse Polish Notation for WHERE/expression evaluation.'],
     ['evalRpn(rpn, row)', 'Evaluates RPN against a row (case-insensitive column lookup; LIKE with % wildcards).'],
     ['parseAlias(f)', 'Splits "expr AS alias".'],
     ['runSql(sql, dataset)', 'Entry point: parses SELECT/WHERE/GROUP BY/ORDER BY/LIMIT, returns rows.'],
     ['project(list, fields)', 'SELECT projection (columns or computed expressions).'],
     ['groupAgg(list, field, fields)', 'GROUP BY with COUNT(*)/AVG/SUM/MIN/MAX.'],
     ['splitTop / colKey / colVal / meanCol / sumCol / extreme / round1', 'Comma-split respecting parens; case-insensitive column helpers; aggregate column math.']],
    widths=[2.6, 6.4])

PB()

# =====================================================================
# 6. POPUP.JS REFERENCE
# =====================================================================
H1('6. Functional Reference - popup.js')
P('popup.js is one IIFE holding STATE and all UI behaviour. It references the engine as E = '
  'window.ApexEngine. Grouped by area below.')

H2('6.1 State and storage shim')
P('STATE holds workbook, sheetName, raw, rows, clean, dropped, currency, pivot (last result for '
  'drill/export), drillRows, sql, aiReport and alerts. The store shim wraps chrome.storage.local when '
  'present and localStorage otherwise, so settings persist whether loaded as an extension or opened as a '
  'plain page.')

H2('6.2 Formatting and DOM helpers')
TBL(['Function', 'Purpose'],
    [['nfIN, money, intf, pct01, pctRaw, rating, numf', 'Indian-grouped number / rupee / integer / percent (fraction & raw) / 2dp rating formatters.'],
     ['fmtByType(v, type)', 'Dispatch formatter by column type (currency/int/pct01/pctRaw/rating/num/text).'],
     ['measureFmtType(fmt)', 'Maps an engine measure fmt to a formatter type.'],
     ['el(tag, cls, txt) / clear(node) / $(id) / emptyMsg(text)', 'Tiny DOM constructors and helpers.'],
     ['escapeHtml(s)', 'HTML-escapes user/data strings before any innerHTML use.']],
    widths=[3.0, 6.0])

H2('6.3 renderTable(wrap, columns, rows, opts) - the universal grid')
P('Every table in the app is built by this one function. A column spec is { key, label, type, get(row), '
  'render(row,val)?, sortable?, colorVsAvg? }. It draws a sortable header (click to toggle asc/desc), '
  'formats cells by type, supports custom render returning a DOM Node or text (text is set via '
  'textContent, never innerHTML, so a product title containing "<" cannot break layout or inject '
  'markup), an optional total row, per-row class, and vs-average colouring. Reused by the leaderboard, '
  'type table, matrix, trends, A-vs-B, SQL results, alerts and the drill drawer.')

H2('6.4 Data loading path')
TBL(['Function', 'Role'],
    [['chooseSheet(workbook)', 'Auto-selects "Consolidated Data" (or the largest sheet).'],
     ['ingestRows(json, sheetName)', 'Shared ingest: normalizeRows -> cleanRows -> set STATE; resets STATE.aiReport/alerts so stale reports are never emailed; calls refreshAll.'],
     ['loadSheet(name)', 'sheet_to_json for a named sheet, then ingestRows. Used by file load and Sheet cycling.'],
     ['handleFile(file)', 'FileReader -> XLSX.read -> chooseSheet -> loadSheet, with error handling.'],
     ['cycleSheet()', 'Cycles to the next sheet in the workbook.'],
     ['window.ApexAnalyzerPro', 'Automation hook exposing { ingestRows, state, engine } for tests/power users.']],
    widths=[2.6, 6.4])

H2('6.5 refreshAll and selector population')
P('refreshAll() updates the dataset stat line and calls every pane renderer plus populateDataDrivenSelects(), '
  'which fills all data-dependent dropdowns (snapshot dates, brands, types, sizes, platforms, '
  'thicknesses, L/B/H) and rebuilds the pivot and trend filter rows.')

H2('6.6 Overview pane')
P('snapshotData(date) filters to one snapshot; overviewData() further applies the Overview Platform and '
  'Product Type filters. renderOverview() draws the KPI tiles, brand leaderboard (with ranking columns), '
  'average-price bars and the Type Master table; aggregateTotals() builds the ALL TYPES total row.')

H2('6.7 Pivot pane')
TBL(['Function', 'Role'],
    [['PIVOT_FILTER_DIMS / pivotFilterState', 'The eight curated filter dimensions and their current values.'],
     ['initPivotControls()', 'Populates the six config selects and wires change handlers, swap, reset, export.'],
     ['buildPivotFilters(...)', 'Builds the filter dropdowns from distinct values.'],
     ['currentPivotConfig()', 'Reads the selects + filters into a buildPivot cfg.'],
     ['renderPivot()', 'Calls buildPivot, computes the heatmap scale, builds multi-level headers, renders cells via valCell(), wires drill clicks, updates the stat line.'],
     ['onDrill(row, cell, isTotalCol)', 'Resolves pivot.drill[cellId] -> record indices -> openDrawer with a descriptive title/subtitle.'],
     ['exportPivot()', 'Flattens the pivot (row labels + leaf columns + total) to a sheet.']],
    widths=[2.8, 6.2])

H2('6.8 Drill-down drawer')
P('openDrawer(title, sub, idxList) maps indices to records and renders the full listing table; '
  'closeDrawer() hides it (also bound to the mask click and Escape). exportDrill() exports the current '
  'slice; flattenRow(r) is the canonical record -> export-row shape used by several exporters.')

H2('6.9 Matrix, Trends, A-vs-B panes')
TBL(['Function', 'Role'],
    [['MATRIX_DIMS / renderMatrixControls / renderMatrix / exportMatrix', 'Eight-axis spec bucket; table with vs-market delta and a product link; export.'],
     ['TREND_FILTER_DIMS / buildTrendFilters / trendFilters', 'Trends scoping filters (platform/type/size/L/B/H).'],
     ['renderTrends()', 'Measure-over-time pivot -> date columns + first-to-last delta.'],
     ['renderTimeTravel() / indexByProduct() / exportTimeTravel()', 'Snapshot A vs B matched by product id; per-product price/rating deltas.'],
     ['initAvsBControls / abOpts / renderAvsB', 'Reads all A-vs-B filters into opts, calls compareAvsB, renders the 12-metric table and the top-brands panel.']],
    widths=[3.4, 5.6])

H2('6.10 SQL, AI, Alerts, Notifications')
TBL(['Function', 'Role'],
    [['SQL_SAMPLES / SQL_SCHEMA / renderSqlSchema', 'Sample queries and the schema reference panel.'],
     ['sqlDataset()', 'Flattens canonical records into the queryable row shape.'],
     ['runSqlQuery() / exportSql()', 'Runs E.runSql, renders results with type inference, exports.'],
     ['runAi() / callOllama() / datasetDigest() / buildStatReport() / showAiReport()', 'AI orchestration: try Ollama, else statistical report; both render into the report pane.'],
     ['alertOpts / sevBadge / renderAlerts', 'Reads thresholds, builds alerts via E.buildAlerts, renders summary chips, action items and the alert table.'],
     ['notify / loadNotifySettings / saveNotifySettings', 'Notification log and persisted settings.'],
     ['htmlToText / alertDigestText / alertDigestHtml', 'Build the email bodies (text + HTML).'],
     ['openMailto / postWebhook / pushNotification', 'Delivery: mailto anchor click or webhook POST (with chrome.permissions.request on the origin, port-stripped, try/catch with mailto fallback).']],
    widths=[3.4, 5.6])

H2('6.11 Init and navigation')
P('initNav() wires the sidebar tab switching. The DOMContentLoaded handler initializes every subsystem, '
  'wires the top-bar/file/drawer/overview/alert/notification controls, renders the SQL schema, and loads '
  'saved notification settings.')

PB()

# =====================================================================
# 7. MARKUP & STYLING
# =====================================================================
H1('7. Markup, Styling & Extension Plumbing')

H2('7.1 popup.html')
P('A static shell: a top bar (logo + wordmark + dataset stats + Load/Sheet buttons), a left sidebar with '
  'seven text-only nav buttons, a content area with seven .pane sections, and a fixed drill-down drawer + '
  'mask. Scripts load in order: xlsx.mini.min.js, engine.js, popup.js. The original Apex double-peak logo '
  'is an inline SVG; the favicon is linked from favicon.ico.')

H2('7.2 popup.css and the no-icon rule')
P('A token-based design system (CSS custom properties for colours, surfaces, borders, radii, fonts). '
  'Components: cards, KPI tiles, the grid table, the pivot table (sticky headers, heatmap cells, frozen '
  'row headers), the drawer, chips, badges, bars and the SQL editor. The no-icon rule is enforced by '
  'construction: dropdown carets are drawn with CSS gradients (geometric, not glyphs); sort indicators '
  'are the words "asc"/"desc"; status uses coloured text/badges. A code-point scan in CI fails the build '
  'on any glyph above U+2100 except the rupee sign.')

H2('7.3 manifest.json and background.js')
P('manifest.json (MV3): permissions storage, unlimitedStorage, downloads; host_permissions for the local '
  'Ollama endpoint; optional_host_permissions for arbitrary https/http so the webhook can request access '
  'at runtime; an action with icons; a background service worker. background.js opens popup.html in a '
  'full tab when the toolbar action is clicked (a popup window is too small for a workbench).')

PB()

# =====================================================================
# 8. DEVOPS: BUILD FROM SCRATCH
# =====================================================================
H1('8. DevOps: Building the Tool From Scratch')

H2('8.0 Build philosophy')
P('Build the engine before the UI, and test the engine against ground truth before building anything on '
  'top of it. Because engine.js is DOM-free, it can be driven headlessly in Node against the real '
  'dataset; if its numbers match the source workbook, every UI surface that consumes it inherits that '
  'correctness. There is no compile/bundle step - the repository is the deployable artifact.')

H2('8.1 Environment setup')
BUL([
    ('Browser:', 'Chrome or any Chromium with chrome://extensions and Developer mode.'),
    ('Editor:', 'any; no language server required (plain JS).'),
    ('Node (tests only):', 'any Node 18+. On this machine there is no global node/npm; a usable Node ships in the Playwright cache at C:\\Users\\DARSHAN\\AppData\\Local\\ms-playwright-go\\<ver>\\node.exe.'),
    ('Python (fixtures/handbook):', 'Python 3 with openpyxl (data fixtures) and python-docx (this handbook). pip is reachable via "python -m pip".'),
    ('No package manager required at runtime:', 'SheetJS is vendored as xlsx.mini.min.js.'),
])

H2('8.2 Phase 1 - scaffold the extension')
P('Deliverables: manifest.json (MV3), background.js (open tab), popup.html shell with the top bar, '
  'sidebar and empty panes, popup.css tokens, and an empty engine.js/popup.js. Exit criteria: the '
  'extension loads unpacked, the workbench tab opens, tabs switch, no console errors.')

H2('8.3 Phase 2 - the pure engine and normalization')
P('Build, test-first: cleanKey/toNum/toISODate, FIELD_ALIASES + buildHeaderMap (two-phase), '
  'normalizeRows, cleanRows, DIMENSIONS/MEASURES/AGGS. Write the Node harness in parallel and assert '
  'against workbook ground truth (counts and averages). Exit criteria: normalized count equals raw '
  'count; snapshot counts and brand merges match the workbook exactly.')

H2('8.4 Phase 3 - pivot engine and drill index')
P('Build aggregate, applyFilters and buildPivot with the drill map. Exit criteria: per-(brand,date) and '
  'per-(type,date) averages/counts match cached workbook pivot values; drill index lengths reconcile to '
  'cell counts; multi-level pivots build with subtotals.')

H2('8.5 Phase 4 - UI controller and renderTable')
P('Build STATE, the storage shim, formatters, el/clear/$ helpers, the universal renderTable, fillSelect, '
  'the data-loading path (handleFile -> ingestRows -> refreshAll), and the automation hook. Exit '
  'criteria: loading a file populates the dataset stat line and the overview without errors.')

H2('8.6 Phase 5 - analytical panes')
P('Overview, Pivot Explorer (with heatmap + drawer), Competitive Matrix, Trends, A vs B. Each is a '
  'render function that calls the engine and feeds renderTable. Exit criteria: every pane renders on the '
  'real data; drill-down opens with reconciling counts; filters narrow correctly.')

H2('8.7 Phase 6 - SQL workspace')
P('Wire the engine SQL parser to an editor, sample dropdown, schema panel, run and export. Exit '
  'criteria: WHERE/GROUP BY/ORDER BY/LIMIT queries return correct rows; GROUP BY counts match the '
  'pivot.')

H2('8.8 Phase 7 - AI, alerts and notifications')
P('Build buildAlerts in the engine (test it), then the Alert Center, the statistical report, the '
  'optional Ollama call, and the email/webhook layer with persisted settings. Exit criteria: alerts are '
  'severity-sorted and reconcile; mailto and webhook payloads build correctly; settings persist.')

H2('8.9 Phase 8 - design polish and no-icon enforcement')
P('Apply the design tokens, add the logo/favicon, and add the emoji/icon code-point scan to the test '
  'gate. Exit criteria: the scan reports clean; the UI is consistent; no console warnings.')

PB()

# =====================================================================
# 9. TESTING & QUALITY
# =====================================================================
H1('9. Testing & Quality Strategy')

H2('9.1 Ground-truth back-testing')
P('The defining quality practice: extract known-correct numbers from the source artifact (the workbook '
  'pivot view, Type Master and A-vs-B sheets carry cached values) and assert the engine reproduces them '
  'to the decimal. Examples used as anchors: snapshot row count 8,087; Sleepwell 4,016 models at avg '
  '15,965.599; Kurlon 2,758; Grid 695 at 19,475.161; % in stock 0.78337 / 0.86439 / 0.98571; total '
  'reviews 657,140. If these pass, the aggregation semantics are correct.')

H2('9.2 The Node back-test harness (test/backtest.js)')
P('Loads a JSON fixture of the real rows, runs normalizeRows -> cleanRows -> buildPivot/groupSummary/'
  'compareAvsB/buildAlerts/runSql, and asserts ~59 invariants: header mapping, snapshot integrity, '
  'Type Master parity, A-vs-B parity, pivot cell values, drill reconciliation, multi-level pivots, the '
  'new L/B/H dimensions and ranking stats, filter exactness, and every regression fix.')
CODE('# fixture is generated from the workbook with openpyxl, then:\n'
     '& "$PlaywrightNode\\node.exe" test\\backtest.js\n'
     '# expected tail:  BACK-TEST RESULT:  59 passed, 0 failed')

H2('9.3 The ingestion test (test/ingest_test.js)')
P('Reads the real .xlsx through the vendored SheetJS build (the exact path the UI uses) and re-checks '
  'the workbook anchors end-to-end - proving the parse + normalize + pivot chain, not just the engine in '
  'isolation.')

H2('9.4 Browser smoke testing')
P('Serve the folder statically and drive the page (navigate, inject a dataset via the automation hook, '
  'click tabs, click a pivot cell, read DOM state, read console). Verifies wiring, rendering, drill-down, '
  'and the absence of console errors - the layer the Node tests cannot reach.')

H2('9.5 Static gates')
BUL([
    ('Syntax check:', 'node --check engine.js popup.js background.js.'),
    ('Manifest validation:', 'parse manifest.json as JSON.'),
    ('Emoji/icon scan:', 'fail on any code point > U+2100 except the rupee sign, across all UI files.'),
])

H2('9.6 Adversarial review')
P('Before a release, run an independent multi-dimension review (engine correctness, UI wiring, '
  'notification/security, requirements coverage) and verify each finding skeptically before fixing. This '
  'practice caught seven real bugs in the enhancement round - including a header-mapping collision and a '
  'filter representation mismatch - that the automated tests alone did not surface. Treat it as a '
  'required gate for non-trivial changes.')

H2('9.7 Quality gate summary')
TBL(['Gate', 'Command / action', 'Pass condition'],
    [['Engine tests', 'node test/backtest.js', '59/59 passed'],
     ['Ingestion test', 'node test/ingest_test.js', '6/6 passed'],
     ['Syntax', 'node --check *.js', 'no errors'],
     ['Manifest', 'JSON parse', 'valid'],
     ['No-icon scan', 'code-point scan', 'clean (ASCII + rupee only)'],
     ['Browser smoke', 'preview + DOM asserts', 'no console errors; panes render'],
     ['Adversarial review', 'multi-agent review', '0 unresolved confirmed findings']],
    widths=[2.0, 3.0, 4.0])

PB()

# =====================================================================
# 10. CI/CD & RELEASE
# =====================================================================
H1('10. CI/CD & Release')

H2('10.1 Suggested pipeline')
P('A minimal CI that runs on every push and blocks merge on any gate failure:')
CODE(
"name: apex-ci\n"
"on: [push, pull_request]\n"
"jobs:\n"
"  verify:\n"
"    runs-on: ubuntu-latest\n"
"    steps:\n"
"      - uses: actions/checkout@v4\n"
"      - uses: actions/setup-node@v4\n"
"        with: { node-version: 20 }\n"
"      - name: Syntax check\n"
"        run: node --check engine.js && node --check popup.js && node --check background.js\n"
"      - name: Validate manifest\n"
"        run: node -e \"JSON.parse(require('fs').readFileSync('manifest.json','utf8'))\"\n"
"      - name: No-icon scan\n"
"        run: node tools/scan_icons.js\n"
"      - name: Engine back-test\n"
"        run: node test/backtest.js          # fixture committed or generated in a prior step\n"
"      - name: Package extension\n"
"        if: startsWith(github.ref, 'refs/tags/v')\n"
"        run: |\n"
"          zip -r apex-analyzer-pro.zip manifest.json background.js popup.* engine.js \\\n"
"              xlsx.mini.min.js icon*.png favicon.ico README.md\n"
"      - uses: actions/upload-artifact@v4\n"
"        if: startsWith(github.ref, 'refs/tags/v')\n"
"        with: { name: extension, path: apex-analyzer-pro.zip }"
)

H2('10.2 Versioning')
P('Use semantic versioning in manifest.json (version) and tag releases vMAJOR.MINOR.PATCH. Bump on every '
  'distributed change - Chrome only updates when the manifest version increases. Keep a CHANGELOG.')

H2('10.3 Packaging')
P('The artifact is a ZIP of the runtime files only. Exclude test/, docs/, the data fixture, and the '
  'handbook so the package stays small and free of sample data.')

H2('10.4 Distribution options')
TBL(['Channel', 'When', 'Notes'],
    [['Load unpacked', 'Dev / single user', 'Fastest; no signing'],
     ['Chrome Web Store (unlisted)', 'Small team', 'Private link; Google review; auto-update'],
     ['Enterprise force-install', 'Org-wide', 'ExtensionInstallForcelist policy + self-hosted .crx or Web Store id'],
     ['Packed .crx', 'Manual share', 'Signed package; users must allow it']],
    widths=[2.6, 2.0, 4.4])

H2('10.5 Release checklist')
NUM([
    'All quality gates green (Section 9.7).',
    'manifest version bumped; CHANGELOG updated.',
    'README and this handbook updated for new features.',
    'Smoke-test the packaged ZIP as a fresh unpacked load.',
    'Tag the release and publish the artifact.',
])

PB()

# =====================================================================
# 11. EXTENSIBILITY PLAYBOOK
# =====================================================================
H1('11. Extensibility Playbook')
P('Common changes and exactly where to make them. The engine\'s catalogue-driven design means most '
  'additions are data, not code.')

H2('11.1 Add a new dimension (group/filter axis)')
NUM([
    'engine.js: add { key, label, get(record) } to DIMENSIONS. Add header aliases to FIELD_ALIASES and a field on the canonical record if it is a new source column.',
    'It now appears automatically in every pivot/trend dimension selector. To expose it as a curated filter, add its key to PIVOT_FILTER_DIMS / MATRIX_DIMS / TREND_FILTER_DIMS in popup.js and to its lookup map.',
    'Add a back-test asserting distinctValues and a pivot over it.',
])

H2('11.2 Add a new measure')
NUM([
    'engine.js: add { key, label, fmt, get(record) } to MEASURES with the right null rule (posOrNull / nonNegOrNull).',
    'It appears automatically in the pivot and trends value selectors. Add a formatter type mapping if fmt is new.',
])

H2('11.3 Add a new alert detector')
NUM([
    'engine.js, inside buildAlerts: compute the condition, push { severity, type, metric, brand, model, title, detail, action }, respect maxPerCategory and a sensible floor.',
    'No UI change needed - renderAlerts and the email digests render any alert shape.',
    'Add a back-test asserting the new alert appears on a crafted/real case and counts reconcile.',
])

H2('11.4 Add a new pane/tab')
NUM([
    'popup.html: add a sidebar nav button (data-pane="x") and a <section class="pane" id="pane-x">.',
    'popup.js: write renderX(), call it from refreshAll(), and wire any controls in DOMContentLoaded.',
    'Reuse renderTable and the engine; do not re-implement aggregation.',
])

H2('11.5 Support a new source-column naming')
P('Add the cleaned alias to the relevant FIELD_ALIASES entry (standardized names first). The two-phase '
  'mapper handles precedence; add a buildHeaderMap regression test for the new header.')

PB()

# =====================================================================
# 12. SECURITY & PRIVACY
# =====================================================================
H1('12. Security, Privacy & Compliance')
BUL([
    ('Local-only by default.', 'All parsing and analysis happen in the browser. No telemetry, no analytics, no remote storage.'),
    ('Two opt-in network paths.', 'The local Ollama call (localhost:11434) and the user-configured outbound webhook. Both are user-initiated; the webhook requests host permission at send time and never fires without a configured URL.'),
    ('XSS hygiene.', 'renderTable sets text via textContent by default; any HTML assembled from data (alert digests, AI report) runs through escapeHtml. Product titles with "<" or "&" cannot inject markup.'),
    ('mailto safety.', 'Recipients, subject and body are percent-encoded with encodeURIComponent, preventing header/body injection through crafted data.'),
    ('Permissions are least-privilege.', 'storage/unlimitedStorage/downloads plus a single fixed Ollama host; broad host access is optional and granted per-origin at runtime only when the user uses the webhook.'),
    ('Data egress warning.', 'A webhook sends your competitive data to a third party. Treat the endpoint as a data-processing boundary; prefer an internally controlled relay.'),
    ('No secrets in code.', 'There are no API keys; the webhook URL and recipients live in local storage, not the codebase.'),
])

PB()

# =====================================================================
# 13. KNOWN LIMITATIONS
# =====================================================================
H1('13. Known Limitations & Technical Debt')
P('A candid inventory - the tool is good, not perfect. Each item names the trade-off and a path forward.')
TBL(['Area', 'Limitation', 'Path forward'],
    [['Type safety', 'Plain JS, no types; refactors rely on tests', 'Add JSDoc + tsc --checkJs, or migrate engine to TS with a build step'],
     ['Persistence', 'Loaded dataset is not saved across sessions (by design)', 'Optional IndexedDB cache with size guard'],
     ['SQL engine', 'Subset only: no JOIN, sub-queries, HAVING, DISTINCT, functions in WHERE', 'Document the grammar (Appendix C); extend incrementally with tests'],
     ['Email', 'mailto length-limited; true unattended send needs a webhook relay', 'Provide a reference Apps Script / serverless relay'],
     ['Pivot depth', 'Max two row and two column levels', 'Generalize the recursion if deeper cross-tabs are needed'],
     ['Performance', 'Single-thread, in-memory; very large files (>~200k rows) can lag rendering', 'Virtualized tables; move aggregation to a Web Worker'],
     ['Heatmap', 'Scale is global per render, not per column', 'Per-column or per-row normalization option'],
     ['DOM tests', 'No automated DOM test; browser checks are scripted/manual', 'Add Playwright e2e once a runner is in CI'],
     ['Alert tuning', 'Thresholds are heuristic defaults', 'Expose a settings profile per category'],
     ['Brand identity', 'Relies on Brand (Std); raw fallback is best-effort casing', 'Maintain an explicit alias map for new brands']],
    widths=[1.7, 3.7, 3.6])

PB()

# =====================================================================
# 14. APPENDICES
# =====================================================================
H1('14. Appendices')

H2('Appendix A - Field alias reference')
P('Canonical field -> cleaned aliases (standardized names first; first match wins via the two-phase '
  'mapper).')
TBL(['Field', 'Aliases (cleaned)'],
    [['scrapeDate', 'scrapedate, rundate, date'],
     ['platform', 'platform, store, marketplace'],
     ['brandStd / brandRaw', 'brandstd, standardbrand / brandname, brand, manufacturer'],
     ['title / url / pid', 'producttitle,title,name / urllink,url,link,href / productidasinfsn,asin,fsn'],
     ['activePrice / effPrice / mrp', 'activeprice,sellingprice,price / effectiveprice / mrpvalue,mrp,originalprice'],
     ['discount / rating / reviews', 'discount,discountpercent / starrating,rating,stars / reviewscount,reviewcount'],
     ['wow / emi / bsr', 'wowbestpricehist,wowprice,bestprice / emiprice,emi / bestsellerrank,bsr,rank'],
     ['productType / size / thickness', 'producttype,typestd / mattresssize,sizestd,size / thicknessin,thickness'],
     ['dim / length / breadth / heightDim', 'dimensionslxbhin,dimensions / lengthin,length / breadthin,breadth,width / heightin,height'],
     ['availability / inStock / prime', 'availability,stockstatus / instockflag,instock / primeassuredflag,assured'],
     ['seller / warranty / recordStatus', 'sellerfulfilledby,fulfilledby,seller / warrantyduration,warranty / recordstatus,status']],
    widths=[2.7, 6.3])

H2('Appendix B - SQL grammar (supported subset)')
CODE("SELECT <field-list | *>\n"
     "FROM ?\n"
     "[ WHERE <expr> ]            -- AND OR LIKE = != < > <= >= + - * / ( )\n"
     "[ GROUP BY <field> ]        -- COUNT(*) AVG(c) SUM(c) MIN(c) MAX(c)\n"
     "[ ORDER BY <field> [ASC|DESC] ]\n"
     "[ LIMIT <n> ]\n"
     "field := column | expr [ AS alias ]\n"
     "LIKE patterns: '%term%', 'term%', '%term'")

H2('Appendix C - Alert severity matrix')
TBL(['Type', 'High', 'Medium', 'Low'],
    [['Undercut', '>= 15%', '>= 7%', '1-6%'],
     ['Rating', '< 3.8 stars', '3.8 to floor', '-'],
     ['Stock', '-', 'out of stock', '-'],
     ['Competitor price drop', '>= 12%', '5-11%', '-'],
     ['Discount gap', '-', '-', '>= 10pt']],
    widths=[2.4, 2.2, 2.2, 2.2])

H2('Appendix D - Glossary')
TBL(['Term', 'Meaning'],
    [['Snapshot', 'All rows for one scrape (crawl) date'],
     ['Anchor brand', 'The brand alerts/reports are written from the perspective of'],
     ['Leaf column', 'A bottom-level data column in the pivot (after column grouping)'],
     ['Drill-down', 'The underlying listings that make up a pivot cell'],
     ['Bucket', 'A specific spec combination (type/size/thickness/dimensions)'],
     ['BSR', 'Bestseller Rank; lower is better'],
     ['Ground truth', 'Known-correct values used to verify the engine']],
    widths=[2.2, 6.8])

H2('Appendix E - Commands quick reference')
CODE("# Run engine back-test (Playwright node on this machine)\n"
     '& "C:\\Users\\DARSHAN\\AppData\\Local\\ms-playwright-go\\1.57.0\\node.exe" test\\backtest.js\n\n'
     "# Run ingestion test against the real workbook\n"
     '& "$node" test\\ingest_test.js\n\n'
     "# Syntax check\n"
     '& "$node" --check engine.js ; & "$node" --check popup.js\n\n'
     "# Regenerate this handbook\n"
     "python docs\\build_handbook.py")

# ---- finalize ----
add_footer()
# Tell Word to refresh fields (the TOC) when the document is opened, so the
# table of contents populates with page numbers without a manual "Update Field".
# python-docx's default settings.xml ships a <w:zoom> with no percent attribute,
# which the strict OOXML schema rejects. Set it explicitly to clear that error.
_zoom = doc.settings.element.find(qn('w:zoom'))
if _zoom is not None and _zoom.get(qn('w:percent')) is None:
    _zoom.set(qn('w:percent'), '100')

import os
try:
    doc.save(OUT)
    saved = OUT
except PermissionError:
    saved = OUT[:-5] + '_fixed.docx'
    doc.save(saved)
print('Saved:', saved)
print('Sections:', sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading 1')))

